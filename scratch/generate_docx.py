
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx(md_path, docx_path):
    doc = docx.Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            p.add_run(line.replace('**', '')).bold = True
        elif line.startswith('|'):
            # This is a very simple table parser for the specific table in the MD
            if '---' in line: continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if not cells: continue
            
            # Check if table exists, else create
            if not hasattr(create_docx, 'current_table'):
                create_docx.current_table = doc.add_table(rows=0, cols=len(cells))
                create_docx.current_table.style = 'Table Grid'
            
            row_cells = create_docx.current_table.add_row().cells
            for i, cell_text in enumerate(cells):
                row_cells[i].text = cell_text.replace('**', '')
        else:
            if hasattr(create_docx, 'current_table'):
                delattr(create_docx, 'current_table')
            
            p = doc.add_paragraph()
            # Handle some basic bolding inside lines
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0:
                    run.bold = True
            
            if line.startswith('* '):
                p.style = 'List Bullet'

    doc.save(docx_path)
    print(f"Successfully saved to {docx_path}")

if __name__ == "__main__":
    create_docx('SRS_Updated.md', 'SRS_Updated.docx')
